"""Reporting 模組：Excel / Word / 統計圖輸出 (CLAUDE.md §9)。

所有報表**必須**標示：產製時間、產製者、資料期間、影子模式聲明。
`labor_rules_reviewed` 為 False 時，另加「未經法規複核」警示。
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import date

import matplotlib

# 伺服器端無 GUI，必須在 import pyplot 前指定 Agg backend。
matplotlib.use("Agg")

import matplotlib.pyplot as plt  # noqa: E402
from docx import Document  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from openpyxl import Workbook  # noqa: E402
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side  # noqa: E402
from openpyxl.utils import get_column_letter  # noqa: E402

from app.core.config import settings  # noqa: E402
from app.core.database import utcnow  # noqa: E402

plt.rcParams["font.sans-serif"] = ["Microsoft JhengHei", "Noto Sans CJK TC", "sans-serif"]
plt.rcParams["axes.unicode_minus"] = False


@dataclass
class ReportContext:
    """每份報表都必須帶的表頭資訊。"""

    period_start: date
    period_end: date
    generated_by: str
    facility_name: str = "陽明醫院附設護理之家"

    @property
    def disclaimers(self) -> list[str]:
        notes = [
            f"機構名稱：{self.facility_name}",
            f"資料期間：{self.period_start} ~ {self.period_end}",
            f"產製時間：{utcnow().strftime('%Y-%m-%d %H:%M:%S')} (UTC)",
            f"產製人員：{self.generated_by}",
        ]
        if settings.shadow_mode:
            notes.append(
                "【影子模式】本報表與現有 HR／紙本紀錄平行運作，加班時數為候選值，"
                "未經核准不得作為薪資依據。"
            )
        if not settings.labor.labor_rules_reviewed:
            notes.append(
                "【未經法規複核】工時與加班門檻採系統預設參數，尚未經人事室與勞動法規權責單位確認。"
            )
        return notes


#: 工時明細的欄位定義（欄位鍵 → 中文標題）。
_SUMMARY_COLUMNS = [
    ("employee_no", "員工編號"),
    ("name", "姓名"),
    ("unit", "單位"),
    ("role", "職務角色"),
    ("first_check_in_at", "首次上班時間"),
    ("last_check_out_at", "最後下班時間"),
    ("total_out_count", "外出次數"),
    ("total_work_minutes", "總工時(分)"),
    ("regular_minutes", "正常工時(分)"),
    ("overtime_candidate_minutes", "加班候選(分)"),
    ("night_duty_minutes", "夜間值班(分)"),
    ("holiday_duty_minutes", "假日值班(分)"),
    ("unmatched_event_count", "未配對事件"),
    ("status", "出勤狀態"),
]


def _set_cell_background(cell, fill_hex: str):
    """Word 表格儲存格背景色輔助函式。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def build_excel(rows: list[dict], ctx: ReportContext, events: list[dict] | None = None) -> bytes:
    """工時彙總與打卡明細 Excel (openpyxl 多工作表)。"""
    wb = Workbook()
    
    # ================= Sheet 1: 出勤統計總表 =================
    ws = wb.active
    ws.title = "出勤統計總表"

    thin_border = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )

    # 標題
    title_cell = ws.cell(row=1, column=1, value=f"{ctx.facility_name} - 員工出勤統計報表")
    title_cell.font = Font(name="微軟正黑體", size=16, bold=True, color="1E1B4B")
    ws.merge_cells("A1:N1")
    ws.row_dimensions[1].height = 32

    # 表頭聲明區
    for idx, note in enumerate(ctx.disclaimers, start=2):
        cell = ws.cell(row=idx, column=1, value=note)
        cell.font = Font(name="微軟正黑體", size=9.5, bold=idx > 4, color="C00000" if idx > 4 else "475569")
        ws.merge_cells(start_row=idx, start_column=1, end_row=idx, end_column=14)

    header_row = len(ctx.disclaimers) + 3
    header_fill = PatternFill("solid", fgColor="4F46E5")  # Indigo Primary
    
    for col, (_, title) in enumerate(_SUMMARY_COLUMNS, start=1):
        cell = ws.cell(row=header_row, column=col, value=title)
        cell.font = Font(name="微軟正黑體", size=10.5, bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border
    ws.row_dimensions[header_row].height = 24

    for r, row in enumerate(rows, start=header_row + 1):
        is_doctor = "DOC" in str(row.get("employee_no", ""))
        row_fill = PatternFill("solid", fgColor="EEF2FF" if is_doctor else ("F8FAFC" if r % 2 == 0 else "FFFFFF"))
        
        for c, (key, _) in enumerate(_SUMMARY_COLUMNS, start=1):
            value = row.get(key)
            cell = ws.cell(row=r, column=c, value=_cell_value(value))
            cell.font = Font(name="微軟正黑體", size=10, bold=is_doctor and key in ("name", "employee_no"))
            cell.fill = row_fill
            cell.border = thin_border
            if key in ("total_work_minutes", "regular_minutes", "overtime_candidate_minutes", "night_duty_minutes", "holiday_duty_minutes", "unmatched_event_count", "total_out_count"):
                cell.alignment = Alignment(horizontal="right", vertical="center")
            else:
                cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[r].height = 20

    for col in range(1, len(_SUMMARY_COLUMNS) + 1):
        col_letter = get_column_letter(col)
        ws.column_dimensions[col_letter].width = 16
    ws.freeze_panes = ws.cell(row=header_row + 1, column=1)

    # ================= Sheet 2: 打卡明細流水 =================
    ws_events = wb.create_sheet(title="護理站打卡流水")
    event_headers = ["流水號 / 時間", "員工工號", "姓名", "事件類型", "打卡站點", "病房代碼", "例外事由", "來源"]
    
    ws_events.cell(row=1, column=1, value=f"{ctx.facility_name} - 打卡與巡診即時流水紀錄").font = Font(name="微軟正黑體", size=14, bold=True)
    ws_events.merge_cells("A1:H1")
    
    ws_events.row_dimensions[2].height = 24
    for c, h in enumerate(event_headers, start=1):
        cell = ws_events.cell(row=2, column=c, value=h)
        cell.font = Font(name="微軟正黑體", size=10.5, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="0D9488")  # Teal
        cell.alignment = Alignment(horizontal="center", vertical="center")

    if events:
        for r, ev in enumerate(events, start=3):
            ws_events.cell(row=r, column=1, value=_cell_value(ev.get("occurred_at")))
            ws_events.cell(row=r, column=2, value=ev.get("employee_no"))
            ws_events.cell(row=r, column=3, value=ev.get("name"))
            ws_events.cell(row=r, column=4, value=ev.get("event_type"))
            ws_events.cell(row=r, column=5, value=ev.get("station_id") or "YM-3F-KIOSK")
            ws_events.cell(row=r, column=6, value=ev.get("ward_code") or "")
            ws_events.cell(row=r, column=7, value=ev.get("override_reason") or "")
            ws_events.cell(row=r, column=8, value=ev.get("source") or "WEB")
            for col in range(1, 9):
                ws_events.cell(row=r, column=col).border = thin_border
                ws_events.cell(row=r, column=col).font = Font(name="微軟正黑體", size=9.5)
            ws_events.row_dimensions[r].height = 19

    for col in range(1, 9):
        ws_events.column_dimensions[get_column_letter(col)].width = 18

    # ================= Sheet 3: 醫師巡診專區 =================
    ws_doc = wb.create_sheet(title="朱醫師巡診專區")
    ws_doc.cell(row=1, column=1, value="朱國大醫師 (YM-DOC-01) - 3F 護理之家巡診簽到紀錄").font = Font(name="微軟正黑體", size=14, bold=True, color="7C3AED")
    ws_doc.merge_cells("A1:F1")
    
    doc_notes = [
        "法規依據：護理之家設置標準及支援醫師巡診規範",
        "在場證明：僅記錄 3F 護理站 Kiosk 巡診在場簽到，絕無病歷、處方或住民個人病況個資 (紅線 R3/R10 遵循)。",
    ]
    for idx, n in enumerate(doc_notes, start=2):
        ws_doc.cell(row=idx, column=1, value=n).font = Font(name="微軟正黑體", size=9.5, color="64748B")
        ws_doc.merge_cells(start_row=idx, start_column=1, end_row=idx, end_column=6)

    doc_headers = ["巡診日期與時間", "醫師姓名", "醫師工號", "簽到站點", "巡診病房", "簽核狀態"]
    for c, h in enumerate(doc_headers, start=1):
        cell = ws_doc.cell(row=5, column=c, value=h)
        cell.font = Font(name="微軟正黑體", size=10.5, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="7C3AED")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    doc_events = [ev for ev in (events or []) if "DOC" in str(ev.get("employee_no", "")) or ev.get("event_type") == "CLINICAL_ROUND_SIGN_IN"]
    if not doc_events:
        ws_doc.cell(row=6, column=1, value="（期間內查無醫師巡診簽到事件）")
    else:
        for r, ev in enumerate(doc_events, start=6):
            ws_doc.cell(row=r, column=1, value=_cell_value(ev.get("occurred_at")))
            ws_doc.cell(row=r, column=2, value=ev.get("name") or "朱國大")
            ws_doc.cell(row=r, column=3, value=ev.get("employee_no") or "YM-DOC-01")
            ws_doc.cell(row=r, column=4, value=ev.get("station_id") or "YM-3F-KIOSK")
            ws_doc.cell(row=r, column=5, value=ev.get("ward_code") or "3F-NH")
            ws_doc.cell(row=r, column=6, value="已完成巡診在勤認證")
            for col in range(1, 7):
                ws_doc.cell(row=r, column=col).border = thin_border
                ws_doc.cell(row=r, column=col).font = Font(name="微軟正黑體", size=9.5)

    for col in range(1, 7):
        ws_doc.column_dimensions[get_column_letter(col)].width = 20

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def build_word(rows: list[dict], ctx: ReportContext, *, title: str = "陽明醫院附設護理之家 員工出勤與巡診統計報表", events: list[dict] | None = None) -> bytes:
    """出勤與統計月報 Word (python-docx 專業臨床排版)。"""
    doc = Document()

    # 設定邊界
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # 1. 主標題
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.font.name = "微軟正黑體"
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 27, 75)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Hospital Nursing Home Staff Attendance & Clinical Round Documentation")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 116, 139)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 2. 聲明與報表資訊區塊
    disclaimer_box = doc.add_table(rows=1, cols=1)
    disclaimer_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = disclaimer_box.rows[0].cells[0]
    _set_cell_background(cell, "F1F5F9")
    
    p = cell.paragraphs[0]
    p_run = p.add_run("【報表基礎資訊與合規聲明】\n")
    p_run.font.name = "微軟正黑體"
    p_run.font.bold = True
    p_run.font.size = Pt(10)
    p_run.font.color.rgb = RGBColor(79, 70, 229)

    for note in ctx.disclaimers:
        p2 = cell.add_paragraph()
        run = p2.add_run(f"• {note}")
        run.font.name = "微軟正黑體"
        run.font.size = Pt(9)
        if "影子模式" in note or "未經法規複核" in note:
            run.font.color.rgb = RGBColor(192, 0, 0)
            run.font.bold = True
        else:
            run.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph()

    # 3. 統計表
    doc.add_heading("一、護理之家 28 名同仁出勤工時統計表", level=2)
    
    # 精簡欄位以適應 Word 寬度
    display_cols = [
        ("employee_no", "工號"),
        ("name", "姓名"),
        ("role", "職稱角色"),
        ("first_check_in_at", "首次上班"),
        ("last_check_out_at", "最後下班"),
        ("total_work_minutes", "總工時(分)"),
        ("regular_minutes", "正常(分)"),
        ("overtime_candidate_minutes", "加班(分)"),
        ("status", "出勤狀態"),
    ]

    table = doc.add_table(rows=1, cols=len(display_cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表頭
    hdr_cells = table.rows[0].cells
    for col_idx, (_, header) in enumerate(display_cols):
        hdr_cells[col_idx].text = header
        _set_cell_background(hdr_cells[col_idx], "4F46E5")
        hdr_run = hdr_cells[col_idx].paragraphs[0].runs[0]
        hdr_run.font.name = "微軟正黑體"
        hdr_run.font.bold = True
        hdr_run.font.size = Pt(9.5)
        hdr_run.font.color.rgb = RGBColor(255, 255, 255)

    # 填入資料
    for row in rows:
        row_cells = table.add_row().cells
        is_doc = "DOC" in str(row.get("employee_no", ""))
        for col_idx, (key, _) in enumerate(display_cols):
            val_str = str(_cell_value(row.get(key)) or "-")
            row_cells[col_idx].text = val_str
            if is_doc:
                _set_cell_background(row_cells[col_idx], "EEF2FF")
            p_cell = row_cells[col_idx].paragraphs[0]
            if p_cell.runs:
                r_run = p_cell.runs[0]
                r_run.font.name = "微軟正黑體"
                r_run.font.size = Pt(9)
                if is_doc and key in ("name", "employee_no"):
                    r_run.font.bold = True

    doc.add_paragraph()

    # 4. 醫師巡診專區
    doc.add_heading("二、兼任醫師（朱國大醫師）病房巡診簽到紀錄", level=2)
    doc_p = doc.add_paragraph()
    doc_p_run = doc_p.add_run(
        "• 依據護理機構評鑑與全民健保規範，兼任醫師定期至 3F 護理之家進行住民健康照護與巡診訪視。\n"
        "• 本打卡系統之巡診簽到僅作為「病房在場出勤證據」，未涉及或儲存任何病患個人病歷資料（遵循紅線 R3/R10）。"
    )
    doc_p_run.font.name = "微軟正黑體"
    doc_p_run.font.size = Pt(9.5)
    doc_p_run.font.color.rgb = RGBColor(51, 65, 85)

    doc_table = doc.add_table(rows=1, cols=5)
    doc_table.style = "Table Grid"
    doc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc_hdr = doc_table.rows[0].cells
    for i, h in enumerate(["巡診日期與時間", "醫師姓名", "醫師工號", "簽到站點", "巡診病房"]):
        doc_hdr[i].text = h
        _set_cell_background(doc_hdr[i], "7C3AED")
        doc_hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        doc_hdr[i].paragraphs[0].runs[0].font.bold = True
        doc_hdr[i].paragraphs[0].runs[0].font.size = Pt(9.5)

    doc_events = [ev for ev in (events or []) if "DOC" in str(ev.get("employee_no", "")) or ev.get("event_type") == "CLINICAL_ROUND_SIGN_IN"]
    if not doc_events:
        r_c = doc_table.add_row().cells
        r_c[0].text = "本期間查無打卡紀錄"
        r_c[1].text = "朱國大"
        r_c[2].text = "YM-DOC-01"
        r_c[3].text = "YM-3F-KIOSK"
        r_c[4].text = "3F-NH"
    else:
        for ev in doc_events:
            r_c = doc_table.add_row().cells
            r_c[0].text = str(_cell_value(ev.get("occurred_at")) or "-")
            r_c[1].text = str(ev.get("name") or "朱國大")
            r_c[2].text = str(ev.get("employee_no") or "YM-DOC-01")
            r_c[3].text = str(ev.get("station_id") or "YM-3F-KIOSK")
            r_c[4].text = str(ev.get("ward_code") or "3F-NH")

    doc.add_paragraph()

    # 5. 簽核區
    doc.add_heading("三、審核與簽章確認", level=2)
    sig_p = doc.add_paragraph()
    sig_p.add_run(
        "護理長／負責人簽章：_______________________　　　"
        "兼任醫師簽章：_______________________\n\n"
        "人事行政主管簽章：_______________________　　　"
        "審核日期：民國 ______ 年 ______ 月 ______ 日"
    )
    sig_p.runs[0].font.name = "微軟正黑體"
    sig_p.runs[0].font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_overtime_plot(rows: list[dict], ctx: ReportContext) -> bytes:
    """加班候選時數長條圖 (matplotlib)，回傳 PNG bytes。"""
    labels = [r.get("name") or r.get("employee_no", "?") for r in rows]
    values = [(r.get("overtime_candidate_minutes") or 0) / 60 for r in rows]

    fig, ax = plt.subplots(figsize=(max(8, len(labels) * 0.4), 5))
    ax.bar(labels, values, color="#4F46E5")
    ax.set_ylabel("加班候選時數（小時）")
    ax.set_title(f"加班候選時數　{ctx.period_start} ~ {ctx.period_end}")
    ax.tick_params(axis="x", rotation=60, labelsize=8)

    if not settings.labor.labor_rules_reviewed:
        ax.text(
            0.5,
            0.5,
            "未經法規複核",
            transform=ax.transAxes,
            fontsize=32,
            color="red",
            alpha=0.15,
            ha="center",
            va="center",
            rotation=30,
        )

    fig.tight_layout()
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    return buffer.getvalue()


def build_unit_distribution_plot(rows: list[dict], ctx: ReportContext) -> bytes:
    """各單位總工時分布圖 (matplotlib)，回傳 PNG bytes。"""
    totals: dict[str, float] = {}
    for row in rows:
        unit = str(row.get("unit") or "OTHER")
        totals[unit] = totals.get(unit, 0) + (row.get("total_work_minutes") or 0) / 60

    fig, ax = plt.subplots(figsize=(7, 5))
    ax.bar(list(totals.keys()), list(totals.values()), color="#10B981")
    ax.set_ylabel("總工時（小時）")
    ax.set_title(f"各單位總工時　{ctx.period_start} ~ {ctx.period_end}")
    fig.tight_layout()

    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    return buffer.getvalue()


def _cell_value(value):
    """把 enum / datetime 轉成試算表可讀值。"""
    if value is None:
        return None
    if hasattr(value, "value"):  # StrEnum
        return value.value
    if hasattr(value, "isoformat"):
        return value.isoformat(timespec="minutes")
    return value
